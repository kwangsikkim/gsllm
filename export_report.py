#!/usr/bin/env python3
import argparse
import json
import re
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, List, Optional


DEFAULT_SUMMARY_DIR = Path(__file__).resolve().parent / "summary_output"
DEFAULT_REPORT_DIR = Path(__file__).resolve().parent / "report_output"


SECTION_ARTIFACT_PATTERNS = [
    r"^\s*#{1,6}\s+",
    r"^\s*(요약|배경/목적|핵심 포인트|결정사항|액션 아이템|리스크|이슈|미해결 질문|다음 단계|화자별 요약|품질 점검)\s*$",
    r"^\s*(verdict|warnings|retry_suggested)\s*:\s*.*$",
]
SECTION_ARTIFACT_RE = re.compile("|".join(SECTION_ARTIFACT_PATTERNS), re.IGNORECASE)

BULLETISH_RE = re.compile(r"^\s*([•\-\*\u2022]|\d+\.|\(\d+\))\s+")
JSON_BLOB_HINT_RE = re.compile(r'^\s*\{[\s\S]*"summary"\s*:\s*', re.IGNORECASE)
SUMMARY_VALUE_RE = re.compile(r'"summary"\s*:\s*"([^"]+)"', re.DOTALL)


def utc_now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def origin_from_summary_filename(name: str) -> str:
    if name.endswith("_summary.json"):
        return name[: -len("_summary.json")]
    return Path(name).stem


def find_latest_dir(root: Path) -> Path:
    if not root.exists():
        raise SystemExit(f"Summary output root not found: {root}")
    candidates = [p for p in root.iterdir() if p.is_dir()]
    if not candidates:
        raise SystemExit(f"No summary folders found in {root}")
    return max(candidates, key=lambda p: p.stat().st_mtime)


def load_summary_json(root: Path, origin: Optional[str]) -> Path:
    if origin:
        target = root / origin / f"{origin}_summary.json"
        if target.exists():
            return target
        raise SystemExit(f"No summary json found in {root / origin}")

    latest = find_latest_dir(root)
    candidates = list(latest.glob("*_summary.json"))
    if candidates:
        return max(candidates, key=lambda p: p.stat().st_mtime)
    raise SystemExit(f"No summary json found in {latest}")


def load_qc_json(summary_path: Path) -> Optional[Path]:
    origin_stem = origin_from_summary_filename(summary_path.name)
    qc_path = summary_path.parent / f"{origin_stem}_qc_report.json"
    return qc_path if qc_path.exists() else None


def _parse_json_blob(text: str) -> Optional[Dict[str, Any]]:
    try:
        return json.loads(text)
    except Exception:
        return None


def _as_text(x: Any) -> str:
    if x is None:
        return ""
    if isinstance(x, str):
        return x
    if isinstance(x, dict):
        if "summary" in x:
            return str(x.get("summary") or "")
        return json.dumps(x, ensure_ascii=False)
    return str(x)


def _clean_multiline_text(text: str) -> str:
    if not isinstance(text, str):
        text = str(text)

    t = text.strip()
    if not t:
        return ""

    if t.startswith("{") and JSON_BLOB_HINT_RE.search(t):
        parsed = _parse_json_blob(t)
        if isinstance(parsed, dict) and "summary" in parsed:
            t = _as_text(parsed.get("summary")).strip()
        else:
            m = SUMMARY_VALUE_RE.search(t)
            if m:
                t = m.group(1).strip()

    lines: List[str] = []
    for raw in t.splitlines():
        line = raw.strip()
        if not line:
            continue
        if SECTION_ARTIFACT_RE.match(line):
            continue
        if "who:" in line and "what:" in line and "due:" in line:
            continue
        lines.append(line)

    if len(lines) >= 10:
        filtered: List[str] = []
        for line in lines:
            if BULLETISH_RE.match(line) and any(k in line for k in ["who:", "what:", "due:"]):
                continue
            if SECTION_ARTIFACT_RE.match(line):
                continue
            filtered.append(line)
        lines = filtered

    return " ".join(lines).strip()


def _clean_dict_values(item: Dict[str, Any]) -> Dict[str, Any]:
    cleaned: Dict[str, Any] = {}
    for key, value in item.items():
        text = _clean_multiline_text(_as_text(value))
        if text:
            cleaned[key] = text
    return cleaned


def _sanitize_list(items: Any, *, aggressive: bool = True) -> List[Any]:
    if items is None:
        return []
    if not isinstance(items, list):
        return []

    cleaned: List[Any] = []
    for item in items:
        if isinstance(item, dict):
            compact = _clean_dict_values(item)
            if compact:
                cleaned.append(compact)
            continue

        if isinstance(item, str):
            text = item.strip()
            if not text:
                continue

            compact = _clean_multiline_text(text)
            if not compact:
                continue

            if aggressive and (
                "핵심 포인트" in text
                or "결정사항" in text
                or "액션 아이템" in text
                or "품질 점검" in text
                or "verdict:" in text
            ):
                if len(compact) < 120 and not SECTION_ARTIFACT_RE.search(compact):
                    cleaned.append(compact)
                continue

            cleaned.append(compact)
            continue

        cleaned.append(item)

    return cleaned


def get_summary_fields(summary_data: Dict[str, Any]) -> Dict[str, Any]:
    summary_obj = summary_data.get("summary", {})
    if not isinstance(summary_obj, dict):
        summary_obj = {}

    summary_text = summary_obj.get("summary") or ""
    if isinstance(summary_text, str):
        stripped = summary_text.strip()
        if stripped.startswith("{") and JSON_BLOB_HINT_RE.search(stripped):
            parsed = _parse_json_blob(stripped)
            if isinstance(parsed, dict):
                summary_obj = {**summary_obj, **parsed}

    summary_clean = _clean_multiline_text(_as_text(summary_obj.get("summary") or ""))
    if summary_clean.startswith("{") and '"summary"' in summary_clean:
        m = SUMMARY_VALUE_RE.search(summary_clean)
        if m:
            summary_clean = m.group(1).strip()
    context_clean = _clean_multiline_text(_as_text(summary_obj.get("context") or ""))

    return {
        "summary": summary_clean,
        "context": context_clean,
        "key_points": _sanitize_list(summary_obj.get("key_points")),
        "decisions": _sanitize_list(summary_obj.get("decisions")),
        "action_items": _sanitize_list(summary_obj.get("action_items")),
        "risks": _sanitize_list(summary_obj.get("risks")),
        "issues": _sanitize_list(summary_obj.get("issues")),
        "open_questions": _sanitize_list(summary_obj.get("open_questions")),
        "next_steps": _sanitize_list(summary_obj.get("next_steps"), aggressive=True),
        "speaker_summary": _sanitize_list(summary_obj.get("speaker_summary")),
    }


def write_docx(output_path: Path, origin: str, summary_fields: Dict[str, Any], qc_data: Optional[Dict[str, Any]]) -> None:
    try:
        from docx import Document
    except Exception as exc:
        raise SystemExit(
            "Missing dependency: python-docx. Install via `pip install python-docx`."
        ) from exc

    doc = Document()
    doc.add_heading(f"요약 보고서 - {origin}", level=1)

    def add_text_section(title: str, content: str) -> None:
        doc.add_heading(title, level=2)
        doc.add_paragraph(content.strip() or "없음")

    def add_list_section(title: str, items: List[Any]) -> None:
        doc.add_heading(title, level=2)
        if not items:
            doc.add_paragraph("없음")
            return
        for item in items:
            if isinstance(item, dict):
                text = ", ".join(f"{k}: {v}" for k, v in item.items() if v is not None)
                doc.add_paragraph(text or "{}", style="List Bullet")
            else:
                doc.add_paragraph(str(item), style="List Bullet")

    def add_table_section(title: str, items: List[Any], key_order: Optional[List[str]] = None) -> None:
        doc.add_heading(title, level=2)
        dict_items = [x for x in items if isinstance(x, dict)]
        if not dict_items:
            doc.add_paragraph("없음")
            return

        keys: List[str] = []
        if key_order:
            keys = key_order[:]
        else:
            for d in dict_items:
                for k in d.keys():
                    if k not in keys:
                        keys.append(k)

        table = doc.add_table(rows=1, cols=max(1, len(keys)))
        hdr = table.rows[0].cells
        for i, k in enumerate(keys):
            hdr[i].text = k

        for d in dict_items:
            row = table.add_row().cells
            for i, k in enumerate(keys):
                row[i].text = str(d.get(k, "") if d.get(k, "") is not None else "")

    add_text_section("요약", summary_fields["summary"])
    if summary_fields["context"]:
        add_text_section("배경/목적", summary_fields["context"])

    add_list_section("핵심 포인트", summary_fields["key_points"])
    add_list_section("결정사항", summary_fields["decisions"])
    add_table_section("액션 아이템", summary_fields["action_items"], key_order=["who", "what", "due"])
    add_list_section("리스크", summary_fields["risks"])
    add_list_section("이슈", summary_fields["issues"])
    add_list_section("미해결 질문", summary_fields["open_questions"])
    add_list_section("다음 단계", summary_fields["next_steps"])
    add_list_section("화자별 요약", summary_fields["speaker_summary"])

    if qc_data:
        doc.add_heading("품질 점검", level=2)
        for key in ["verdict", "warnings", "retry_suggested"]:
            if key in qc_data:
                doc.add_paragraph(f"{key}: {qc_data[key]}")

    doc.save(str(output_path))


def write_xlsx(output_path: Path, origin: str, summary_fields: Dict[str, Any], qc_data: Optional[Dict[str, Any]]) -> None:
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Alignment, Font
    except Exception as exc:
        raise SystemExit(
            "Missing dependency: openpyxl. Install via `pip install openpyxl`."
        ) from exc

    wb = Workbook()
    ws = wb.active
    ws.title = "Summary"

    header_font = Font(bold=True)
    wrap = Alignment(wrap_text=True, vertical="top")

    ws.append(["Origin", origin])
    ws.append(["Created At", utc_now_iso()])
    ws.append([])

    ws.append(["Summary"])
    ws["A4"].font = header_font
    ws.append([summary_fields["summary"] or "없음"])
    ws["A5"].alignment = wrap
    ws.append([])

    def write_list(title: str, items: List[Any]) -> None:
        ws.append([title])
        ws.cell(row=ws.max_row, column=1).font = header_font
        if not items:
            ws.append(["없음"])
            ws.append([])
            return
        for item in items:
            if isinstance(item, dict):
                text = ", ".join(f"{k}: {v}" for k, v in item.items() if v is not None)
                ws.append([text or "없음"])
            else:
                ws.append([str(item)])
        ws.append([])

    if summary_fields["context"]:
        ws.append(["배경/목적"])
        ws.cell(row=ws.max_row, column=1).font = header_font
        ws.append([summary_fields["context"]])
        ws.cell(row=ws.max_row, column=1).alignment = wrap
        ws.append([])

    write_list("핵심 포인트", summary_fields["key_points"])
    write_list("결정사항", summary_fields["decisions"])
    write_list("액션 아이템", summary_fields["action_items"])
    write_list("리스크", summary_fields["risks"])
    write_list("이슈", summary_fields["issues"])
    write_list("미해결 질문", summary_fields["open_questions"])
    write_list("다음 단계", summary_fields["next_steps"])
    write_list("화자별 요약", summary_fields["speaker_summary"])

    if qc_data:
        ws.append(["QC Verdict", qc_data.get("verdict", "")])
        ws.append(["QC Warnings", json.dumps(qc_data.get("warnings", []), ensure_ascii=False)])
        ws.append(["QC Retry Suggested", qc_data.get("retry_suggested", "")])

    for row in ws.iter_rows():
        for cell in row:
            cell.alignment = wrap

    wb.save(str(output_path))


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Export report from summary.json + qc_report.json to DOCX/XLSX.")
    parser.add_argument("origin", nargs="?", default=None)
    parser.add_argument("--summary-root", type=str, default=None)
    parser.add_argument("--report-root", type=str, default=None)
    parser.add_argument("--docx", action="store_true")
    parser.add_argument("--xlsx", action="store_true")

    # ✅ 추가: stt_run.py에서 넘기는 --title 지원
    parser.add_argument(
        "--title",
        type=str,
        default=None,
        help="Report filename base (without extension). If omitted, uses origin stem.",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    summary_root = Path(args.summary_root) if args.summary_root else DEFAULT_SUMMARY_DIR

    summary_path = load_summary_json(summary_root, args.origin)
    summary_data = json.loads(summary_path.read_text(encoding="utf-8"))

    qc_path = load_qc_json(summary_path)
    qc_data = json.loads(qc_path.read_text(encoding="utf-8")) if qc_path else None

    summary_fields = get_summary_fields(summary_data)
    origin_stem = origin_from_summary_filename(summary_path.name)

    report_root = Path(args.report_root) if args.report_root else DEFAULT_REPORT_DIR
    output_dir = report_root / origin_stem
    output_dir.mkdir(parents=True, exist_ok=True)

    export_docx = args.docx or not args.xlsx
    export_xlsx = args.xlsx or not args.docx

    # ✅ 추가: 파일명 접두어를 title로 지정 가능
    report_base = args.title.strip() if isinstance(args.title, str) and args.title.strip() else origin_stem

    if export_docx:
        docx_path = output_dir / f"{report_base}.docx"
        write_docx(docx_path, origin_stem, summary_fields, qc_data)
        print(f"Saved: {docx_path}")

    if export_xlsx:
        xlsx_path = output_dir / f"{report_base}.xlsx"
        write_xlsx(xlsx_path, origin_stem, summary_fields, qc_data)
        print(f"Saved: {xlsx_path}")


if __name__ == "__main__":
    main()
